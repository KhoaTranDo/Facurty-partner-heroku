# package word
import docx
import json
# from pymongo import MongoClient
import os
import boto3
import sys
import io
# Open file word

# Get file by url
client = boto3.client('s3')
obj = client.get_object(
    Bucket = 'farcurtypartner',
    Key = sys.argv[1]
)
# Convert to docx
object_as_streaming_body = obj['Body']
object_as_bytes = object_as_streaming_body.read()
object_as_file_like = io.BytesIO(object_as_bytes)

# Get data
doc = docx.Document(object_as_file_like)
if doc.paragraphs[-1].text!="#End" :
    doc.add_paragraph("#End")
    
all_text = doc.paragraphs

# Lưu câu hỏi
answer=[]
# Lưu vị trí
index=0
luutamthoi = {}
answerthamthoi = []
trueanswer=[]
exams=[]
sothutu=[]
so=0
image=''
numimage=1
luuanh=[]
exams=[]
saveimage=[]
checkimage={}


def TimCauHoi(a):
    cauhoi = a.text
    # Nhận câu hỏi và đáp án theo ký tự
    words=['^-^']
    words1 = ['A', 'B', 'C', 'D', 'E', 'F','a','b','c','d','e','f']
    words11 = []
    # Xử lý ký tự đầu đáp án
    for i in words1:
        words11.append(i + ')')
        words11.append(i + '/')
        words11.append(i + '.')
    # XXử lý ký tự đầu câu hỏi
    for i in range(200,-1,-1):
        words.append(str(i)+')')
        words.append('Câu' + str(i))
        words.append('Câu ' + str(i))
        words.append(str(i) + '/')
        words.append(str(i) + '.')
    # Xử lý dòng văn bản chuyển sang ký tự đặt biệt để nhận biết câu hỏi và đáp án
    # For tìm câu hỏi
    for word in words:
       #  For tìm đáp án
       for worda in words11:
            # Kiểm tra có phải là câu hỏi không
            if word in cauhoi[:7].strip():
                cauhoi = cauhoi[:7].replace(word, '^-^',2)+cauhoi[7:]
                a.text=cauhoi
            # Kiểm tra có phải là đáp án không
            if worda in cauhoi[:7].strip():
                for run in a.runs:
                    if run.bold:
                        cauhoi = cauhoi[:7].replace(worda, '$$$-',1)+cauhoi[7:]
                        break
                    else:
                        cauhoi = cauhoi[:7].replace(worda, '/-/',1)+cauhoi[7:]
                        break
                a.text = cauhoi

    return (a.text)


# # Lay so luong anh
# for s in doc.inline_shapes:
#     print(s.height.cm, s.width.cm, s._inline.docPr.id)
# print(len(doc.inline_shapes))

# Lấy hinh anh
import xml.etree.ElementTree as ET
def hasImage(par):
    """get all of the images in a paragraph
    :param par: a paragraph object from docx
    :return: a list of r:embed
    """
    ids = []
    root = ET.fromstring(par._p.xml)
    namespace = {
             'a':"http://schemas.openxmlformats.org/drawingml/2006/main", \
             'r':"http://schemas.openxmlformats.org/officeDocument/2006/relationships", \
             'wp':"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}

    inlines = root.findall('.//wp:inline',namespace)
    for inline in inlines:
        imgs = inline.findall('.//a:blip', namespace)
        for img in imgs:
            id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
        ids.append(id)
    return ids


for a in all_text:
    if ("^-^") in TimCauHoi(a):
        if answerthamthoi !=[]:
            luutamthoi["Answer"] = answerthamthoi
            luutamthoi["Trueanswer"] = trueanswer
            exams.append(luutamthoi)
            luuanh=[]
            luutamthoi={}
            answerthamthoi=[]
            trueanswer=[]
            index = index + 1
        cauhoi = TimCauHoi(a).replace('^-^', '').strip()
        luutamthoi['Question'] = cauhoi
    if ("/-/"or"$$$-") in TimCauHoi(a):
        dapan=TimCauHoi(a).replace('/-/','').strip()
        answerthamthoi.append(dapan)
    if "$$$-" in TimCauHoi(a):
        dapandung = TimCauHoi(a).replace('$$$-', '').strip()
        answerthamthoi.append(dapandung)
        trueanswer.append(dapandung)
    if "#End" in TimCauHoi(a):
        luutamthoi["Answer"] = answerthamthoi
        luutamthoi["Trueanswer"] = trueanswer
        exams.append(luutamthoi)
        break
    if hasImage(a) != []:
        if hasImage(a) not in saveimage:
            saveimage.append(hasImage(a))
            image = 'image' + str(numimage) + '.jgeg'
            checkimage[str(hasImage(a))]=image
            numimage = numimage + 1
            luuanh.append(image)
        else:
            luuanh.append( checkimage[str(hasImage(a))])

    # print(a.text)

print(json.dumps(exams))

# import docx2txt as d2t


# def extract_images_from_docx(path_to_file, images_folder, get_text=False):
#     text = d2t.process(path_to_file, images_folder)
#     if (get_text):
#         return text

# path_to_file = 'public/'+str(sys.argv[1])
# if not os.path.exists('public/images/'+str(sys.argv[1])):
#      os.makedirs('public/images/'+str(sys.argv[1])+'/')
#      images_folder ='public/images/'+str(sys.argv[1]+'/')
# else:
#      images_folder ='public/images/'+str(sys.argv[1]+'/')

# extract_images_from_docx(path_to_file, images_folder)
# extract_images_from_docx(path_to_file, images_folder, get_text=True)

